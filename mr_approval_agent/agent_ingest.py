#!/usr/bin/env python3
"""
agent_ingest.py — Document Ingestion & Field Extraction Pipeline
=================================================================
Stage 1 of the MR Equipment Approval Agent system.

Responsibilities:
  1. Scan datasource/ (a_requirements / b_competitors / c_compliance / d_operations)
  2. Parse every document via MinerU CLI (OCR + table extraction)
  3. Run QwenFieldExtractor on each parsed result to normalise key fields
  4. Build project_snapshot.json — the single structured input for agent_approval.py

Workflow nodes (LangGraph):
  init → parse_documents → extract_fields → build_snapshot → export_log

Reference: supercare/data_agent_task1.py  (MinerUWorkbookParser, QwenExtractorAgent)

Usage:
  python agent_ingest.py --datasource ./datasource --output-root ./outputs/ingest
  python agent_ingest.py --parser-mode openpyxl   # ablation / quick baseline
"""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, TypedDict

from langgraph.graph import END, START, StateGraph
from openai import OpenAI
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

try:
    from json_repair import repair_json
except ImportError:
    repair_json = None

# ── Default paths (override via CLI args) ────────────────────────────────────
DEFAULT_DATASOURCE  = "/srv/mr_approval_agent/datasource"
DEFAULT_OUTPUT_ROOT = "/srv/mr_approval_agent/outputs/ingest"
DEFAULT_CONFIG_PATH = "/srv/mr_approval_agent/config/agent_config.json"

# ── Dual-track datasource layout ─────────────────────────────────────────────
# Track A  ─  静态知识底座 (agent_corpus)
#   Historical / vendor-supplied data; treated as baseline reference.
#   Physical path: <datasource>/agent_corpus/{a_requirements,b_competitors,…}
CORPUS_ROOT = "agent_corpus"

# Track B  ─  动态用户上传 (user_uploads)
#   Authoritative data submitted for *this* approval run.
#   Takes priority over corpus when the same metric is found in both tracks.
#   Physical path: <datasource>/user_uploads/{01_requirements,…}
UPLOADS_ROOT = "user_uploads"

# Logical category names — canonical keys used throughout the pipeline
CORPUS_CATEGORY_DIRS: Dict[str, str] = {
    "a_requirements": "场景边界与立项基础",
    "b_competitors":  "设备选型与竞品资料",
    "c_compliance":   "合规与制度资料",
    "d_operations":   "收益与运营数据",
    "e_systems":      "系统集成与环境",
    "f_validation":   "验收与校验",
}

# Maps user_uploads sub-folder name → logical category key
UPLOAD_DIR_TO_CATEGORY: Dict[str, str] = {
    "01_requirements": "a_requirements",
    "02_competitors":  "b_competitors",
    "03_compliance":   "c_compliance",
    "04_operations":   "d_operations",
    # ── 新增：前端六分类对应目录 ──
    "05_budget":       "a_requirements",   # 预算清单 → 归入立项基础
    "06_price":        "b_competitors",    # 价格依据证明 → 归入设备选型
}

# Legacy alias — keeps code that references CATEGORY_DIRS working unchanged
CATEGORY_DIRS: Dict[str, str] = CORPUS_CATEGORY_DIRS

# ── Fields to extract per category ───────────────────────────────────────────
EXTRACT_FIELDS: Dict[str, List[str]] = {
    # FIX-1: add "设备类型" so CT/PET docs can be flagged and excluded from
    #         device_name population (e.g. 50万以上大型医疗设备申请表 is a CT form).
    # FIX-2: add competitor fields — 科室设备购置认证样表(Hero) embeds a 4-brand
    #         parameter table that belongs logically to b_competitors.
    # FIX-3: add revenue fields — 设备申购报告表2022版 contains actual billing
    #         data (460元/次, 1320人次/月, 3-4yr payback) needed by Agent 4.
    "a_requirements": [
        "设备名称", "设备类型",          # FIX-1: 设备类型 used to detect non-MR docs
        "申请科室", "申请数量", "申请金额",
        "业务需求描述", "临床痛点",
        "竞品品牌列表", "竞品参数对比",   # FIX-2: cross-category competitor data
        "单次收费标准", "月检查量",        # FIX-3: revenue params for Agent 4
        "投入回收期",
    ],
    "b_competitors":  ["品牌", "型号", "磁场强度", "主要技术参数", "含税报价", "使用年限"],
    "c_compliance":   ["注册证号", "设备名称", "生产企业", "有效期", "许可类别"],
    "d_operations":   ["医院名称", "单台日均服务量", "工作饱和度", "维修类型", "维修费用", "预算金额"],
}

# ── Keywords that indicate a document is NOT an MR device application ─────────
# FIX-1: documents matching these keywords will be excluded from device_name
#         population in node_build_snapshot, but still kept in the snapshot
#         data for format/structure reference.
NON_MR_KEYWORDS: List[str] = [
    "体层摄影", "CT", "X射线计算机", "超声", "PET", "核医学",
    "内镜", "腔镜", "手术机器人",
]

# B类：标注内部流转、未获广告审查的宣传材料 — 不得写入对外立项文本（Agent2/6 排除）
INTERNAL_USE_KEYWORDS: List[str] = [
    "仅供内部使用",
    "未获广审",
    "尚未获得广审",
    "内部资料",
]

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".pptx"}


# ─────────────────────────────────────────────────────────────────────────────
# Helpers — device-type guard  (FIX-1)
# ─────────────────────────────────────────────────────────────────────────────
def _is_non_mr_document(fields: Dict[str, Any]) -> bool:
    """
    Return True when extracted fields indicate this document describes a
    non-MR device (e.g. CT, PET, ultrasound).  Used in node_build_snapshot
    to skip such docs when populating the canonical device_name / budget fields.
    """
    device_type = str(fields.get("设备类型") or "")
    device_name = str(fields.get("设备名称") or "")
    combined = device_type + device_name
    return any(kw in combined for kw in NON_MR_KEYWORDS)


def _is_internal_use_item(item: Dict[str, Any]) -> bool:
    """True when filename / summary / extracted fields indicate non-publishable promo."""
    parts: List[str] = [
        str(item.get("file") or ""),
        str(item.get("summary") or ""),
    ]
    fields = item.get("fields") or {}
    for v in fields.values():
        if v:
            parts.append(str(v))
    combined = "\n".join(parts)
    return any(kw in combined for kw in INTERNAL_USE_KEYWORDS)


# ─────────────────────────────────────────────────────────────────────────────
# ConflictableField — dual-track priority & conflict recording
# ─────────────────────────────────────────────────────────────────────────────
#
# Schema of a ConflictableField object:
# {
#   "value":        <any>   — resolved (winning) value; user_upload always wins
#   "origin_type":  str     — "user_upload" | "agent_corpus"
#   "source_file":  str     — filename that provided the resolved value
#   "doc_category": str     — logical category key, e.g. "a_requirements"
#   "conflict":     bool    — True when ≥2 numeric values found across tracks
#   "all_values":   list[{  — every observation, for audit / human review
#                     "value", "origin_type", "source_file", "doc_category"
#                   }]
# }

def _make_conflictable_field(
    value:        Any,
    origin_type:  str,
    source_file:  str,
    doc_category: str,
) -> Dict[str, Any]:
    """Wrap a scalar value in the ConflictableField envelope."""
    entry = {
        "value":        value,
        "origin_type":  origin_type,
        "source_file":  source_file,
        "doc_category": doc_category,
    }
    return {
        "value":        value,
        "origin_type":  origin_type,
        "source_file":  source_file,
        "doc_category": doc_category,
        "conflict":     False,
        "all_values":   [entry],
    }


def _merge_conflictable_field(
    existing:     Dict[str, Any],
    value:        Any,
    origin_type:  str,
    source_file:  str,
    doc_category: str,
) -> Dict[str, Any]:
    """
    Merge a new observation into an existing ConflictableField in-place.

    Priority rules
    ──────────────
    • user_upload always overrides agent_corpus for the resolved .value.
    • If the incoming value differs numerically from the current resolved value,
      conflict is set to True regardless of which track "wins".
    • Two observations from the same track with different numeric values also
      set conflict = True.
    Returns the mutated existing dict.
    """
    def _num(v: Any) -> Optional[float]:
        m = re.search(r"(\d+(?:\.\d+)?)", str(v))
        return float(m.group(1)) if m else None

    def _values_differ(a: Any, b: Any) -> bool:
        na, nb = _num(a), _num(b)
        if na is not None and nb is not None:
            return na != nb
        return str(a).strip() != str(b).strip()

    new_entry = {
        "value":        value,
        "origin_type":  origin_type,
        "source_file":  source_file,
        "doc_category": doc_category,
    }
    existing["all_values"].append(new_entry)

    cur_val  = existing["value"]
    cur_orig = existing["origin_type"]

    if origin_type == "user_upload" and cur_orig == "agent_corpus":
        # user_upload takes the throne
        if _values_differ(value, cur_val):
            existing["conflict"] = True
        existing["value"]        = value
        existing["origin_type"]  = origin_type
        existing["source_file"]  = source_file
        existing["doc_category"] = doc_category

    elif origin_type == cur_orig:
        # Same track — second occurrence; flag conflict if values differ
        if _values_differ(value, cur_val):
            existing["conflict"] = True

    else:
        # New is agent_corpus, existing is user_upload — keep existing, still record conflict
        if _values_differ(value, cur_val):
            existing["conflict"] = True

    return existing


def _tesla_from_magnetic_strength(val: Any) -> Optional[float]:
    """Parse a tesla value from the 磁场强度 field, if present."""
    if val is None or val == "":
        return None
    s = str(val).strip()
    m = re.search(r"(\d+(?:\.\d+)?)\s*[Tt]", s)
    if m:
        return float(m.group(1))
    return None


def _item_is_low_field_15t(
    fields: Dict[str, Any],
    file_name: str,
    summary: str,
) -> bool:
    """
    True when the document clearly describes a ~1.5T (low-field) MR vs the project 3.0T target.
    Trusts explicit 磁场强度 when present; otherwise scans filename, summary, 主要技术参数.
    """
    t_main = _tesla_from_magnetic_strength(fields.get("磁场强度"))
    if t_main is not None:
        if 1.0 <= t_main < 2.0:
            return True
        if t_main >= 2.5:
            return False

    blob = "\n".join(
        [
            file_name,
            str(summary or ""),
            str(fields.get("主要技术参数") or ""),
            str(fields.get("型号") or ""),
        ]
    )
    if re.search(r"1\.5\s*[Tt]|1\.5特斯拉", blob):
        if t_main is not None and t_main >= 2.5:
            return False
        return True
    return False


# ─────────────────────────────────────────────────────────────────────────────
# Config
# ─────────────────────────────────────────────────────────────────────────────
@dataclass
class AppConfig:
    mineru_api_url:     str
    mineru_api_token:   str
    mineru_backend:     str
    mineru_method:      str
    dashscope_api_key:  str
    qwen_base_url:      str
    qwen_model:         str
    qwen_max_tokens:    int
    qwen_temperature:   float
    parser_mode:        str
    max_workers:        int


def load_config(config_path: Path, parser_mode: str = "mineru",
                max_workers: int = 4) -> AppConfig:
    raw = json.loads(config_path.read_text(encoding="utf-8"))
    m, q, r = raw.get("mineru", {}), raw.get("qwen", {}), raw.get("runtime", {})

    mineru_token = os.getenv("MINERU_API_TOKEN", "").strip()
    mineru_url_env = os.getenv("MINERU_API_URL", "").strip()

    # Qwen/DashScope key：优先读环境变量，再 fallback 到 config 文件
    dashscope_key = (
        os.getenv("DASHSCOPE_API_KEY", "").strip()
        or str(q.get("api_key", "") or "").strip()
    )

    api_url = mineru_url_env or str(m.get("api_url", "") or "").strip()
    if api_url and not api_url.startswith(("http://", "https://")):
        api_url = ""

    backend = str(m.get("backend", "pipeline"))

    if backend == "api" and not mineru_token:
        raise ValueError(
            "MinerU 已配置为 backend=api，但未读取到环境变量 MINERU_API_TOKEN。"
            "请执行：export MINERU_API_TOKEN='你的 token'（可先 source setup_env.sh），"
            "或把 config/agent_config.json 里 mineru.backend 改为 pipeline 使用本机 MinerU。"
        )

    return AppConfig(
        mineru_api_url=api_url,
        mineru_api_token=mineru_token,
        mineru_backend=backend,
        mineru_method=str(m.get("method", "auto")),
        dashscope_api_key=dashscope_key,
        qwen_base_url=q.get("base_url", "https://dashscope.aliyuncs.com/compatible-mode/v1"),
        qwen_model=q.get("model", "qwen3-max"),
        qwen_max_tokens=int(q.get("max_tokens", 2000)),
        qwen_temperature=float(q.get("temperature", 0.2)),
        parser_mode=parser_mode or r.get("parser_mode", "mineru"),
        max_workers=max_workers or int(r.get("max_workers", 4)),
    )


# ─────────────────────────────────────────────────────────────────────────────
# Logger
# ─────────────────────────────────────────────────────────────────────────────
class IngestLogger:
    def __init__(self, log_path: Path) -> None:
        self.log_path = log_path
        log_path.parent.mkdir(parents=True, exist_ok=True)
        log_path.write_text("", encoding="utf-8")

    def log(self, level: str, stage: str, message: str,
            payload: Optional[Dict] = None) -> None:
        entry = {
            "ts":      datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "level":   level,
            "stage":   stage,
            "message": message,
            "payload": payload or {},
        }
        with self.log_path.open("a", encoding="utf-8") as fh:
            fh.write(json.dumps(entry, ensure_ascii=False) + "\n")

    def info(self, stage: str, msg: str, payload: Optional[Dict] = None) -> None:
        self.log("INFO",  stage, msg, payload)

    def warn(self, stage: str, msg: str, payload: Optional[Dict] = None) -> None:
        self.log("WARN",  stage, msg, payload)

    def error(self, stage: str, msg: str, payload: Optional[Dict] = None) -> None:
        self.log("ERROR", stage, msg, payload)


# ─────────────────────────────────────────────────────────────────────────────
# MinerU Document Parser  (A1 agent)
# ─────────────────────────────────────────────────────────────────────────────
class MinerUDocumentParser:
    """
    Wraps the mineru CLI.  Falls back to openpyxl baseline for xlsx-only runs.
    Reference: supercare MinerUWorkbookParser._run_mineru_cli / parse_excel
    """

    def __init__(self, config: AppConfig, logger: IngestLogger) -> None:
        self.config = config
        self.logger = logger

    def parse(self, file_path: Path, output_dir: Path) -> Dict[str, Any]:
        suffix = file_path.suffix.lower()
        output_dir.mkdir(parents=True, exist_ok=True)

        if self.config.parser_mode == "openpyxl":
            if suffix == ".xlsx":
                return self._parse_xlsx_baseline(file_path)
            self.logger.warn("mineru_parser",
                             f"openpyxl mode skips non-xlsx: {file_path.name}")
            return {"file": file_path.name, "content": [], "skipped": True}

        if suffix not in SUPPORTED_EXTENSIONS:
            self.logger.warn("mineru_parser", f"Unsupported extension: {file_path.name}")
            return {"file": file_path.name, "content": [], "skipped": True}

        return self._parse_via_mineru(file_path, output_dir)

    def _parse_via_mineru(self, file_path: Path, output_dir: Path) -> Dict[str, Any]:
        # 新版 MinerU CLI 不再支持 backend=api；合法值：
        # pipeline / vlm-http-client / hybrid-http-client / vlm-auto-engine / hybrid-auto-engine
        _VALID_BACKENDS = {
            "pipeline", "vlm-http-client", "hybrid-http-client",
            "vlm-auto-engine", "hybrid-auto-engine",
        }
        backend = self.config.mineru_backend
        if backend not in _VALID_BACKENDS:
            backend = "pipeline"   # api 等已废弃 backend 降级为 pipeline

        cmd = ["mineru", "-p", str(file_path),
               "-o", str(output_dir),
               "-b", backend,
               "-m", self.config.mineru_method,
               "-l", "ch"]           # 指定中文，提升 OCR 精度

        # --api-url 仅在需要连接远程 mineru-api service 时使用
        if self.config.mineru_api_url and backend in ("vlm-http-client", "hybrid-http-client"):
            cmd += ["--api-url", self.config.mineru_api_url]

        self.logger.info("mineru_parser", f"Parsing: {file_path.name}",
                         {"cmd": " ".join(cmd)})
        proc = subprocess.run(cmd, capture_output=True, text=True, check=False)

        if proc.returncode != 0:
            self.logger.error("mineru_parser", f"MinerU failed: {file_path.name}",
                              {"stderr": proc.stderr[-300:]})
            # ── pdfplumber 降级：提取文字层内容 ──────────────────────────────
            self.logger.info("mineru_parser",
                             f"Falling back to pdfplumber: {file_path.name}")
            return self._parse_via_pdfplumber(file_path)

        # Locate content_list.json (office or auto sub-dir)
        for sub in ("office", "auto"):
            candidate = (output_dir / file_path.stem / sub
                         / f"{file_path.stem}_content_list.json")
            if candidate.exists():
                items = json.loads(candidate.read_text(encoding="utf-8"))
                self.logger.info("mineru_parser", f"Done: {file_path.name}",
                                 {"items": len(items)})
                return {"file": file_path.name, "content": items,
                        "content_list_path": str(candidate)}

        self.logger.warn("mineru_parser",
                         f"content_list.json not found after MinerU; trying pdfplumber: {file_path.name}")
        return self._parse_via_pdfplumber(file_path)

    def _parse_via_pdfplumber(self, file_path: Path) -> Dict[str, Any]:
        """pdfplumber 降级解析：提取文字层和表格（不做 OCR）。"""
        suffix = file_path.suffix.lower()
        try:
            if suffix in (".pdf",):
                import pdfplumber
                content: List[Dict] = []
                with pdfplumber.open(str(file_path)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text() or ""
                        if text.strip():
                            content.append({"type": "text", "text": text.strip()})
                        for tbl in (page.extract_tables() or []):
                            if tbl:
                                rows = [
                                    [str(c) if c is not None else "" for c in row]
                                    for row in tbl
                                ]
                                content.append({
                                    "type": "table",
                                    "table_body": rows,
                                })
                self.logger.info("mineru_parser",
                                 f"pdfplumber fallback done: {file_path.name}",
                                 {"blocks": len(content)})
                return {"file": file_path.name, "content": content,
                        "parser": "pdfplumber_fallback"}
            # docx
            elif suffix in (".docx",):
                from docx import Document as DocxDocument
                content = []
                doc = DocxDocument(str(file_path))
                for para in doc.paragraphs:
                    if para.text.strip():
                        content.append({"type": "text", "text": para.text.strip()})
                for table in doc.tables:
                    rows = [[cell.text for cell in row.cells] for row in table.rows]
                    if rows:
                        content.append({"type": "table", "table_body": rows})
                return {"file": file_path.name, "content": content,
                        "parser": "python_docx_fallback"}
            else:
                return {"file": file_path.name, "content": [],
                        "skipped": True, "error": "unsupported for fallback"}
        except Exception as exc:
            self.logger.error("mineru_parser",
                              f"pdfplumber fallback failed: {file_path.name}",
                              {"error": str(exc)})
            return {"file": file_path.name, "content": [], "error": str(exc)}

    def _parse_xlsx_baseline(self, file_path: Path) -> Dict[str, Any]:
        from openpyxl import load_workbook
        wb = load_workbook(str(file_path), data_only=True)
        sheets: Dict[str, List[Dict[str, str]]] = {}
        for sn in wb.sheetnames:
            ws = wb[sn]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            headers = [str(h) if h else f"col_{i}" for i, h in enumerate(rows[0])]
            sheets[sn] = [
                {h: str(v) for h, v in zip(headers, row) if v is not None}
                for row in rows[1:]
            ]
        return {"file": file_path.name, "content": sheets, "parser": "openpyxl"}


# ─────────────────────────────────────────────────────────────────────────────
# Qwen Field Extractor  (A2 agent)
# ─────────────────────────────────────────────────────────────────────────────
class QwenFieldExtractor:
    """
    Calls qwen3-max to normalise raw parsed content into structured fields.
    Reference: supercare QwenExtractorAgent
    """

    def __init__(self, config: AppConfig, logger: IngestLogger) -> None:
        self.config = config
        self.logger = logger
        self.client = OpenAI(api_key=config.dashscope_api_key,
                             base_url=config.qwen_base_url)

    def extract(self, category: str, parsed: Dict[str, Any]) -> Dict[str, Any]:
        target_fields = EXTRACT_FIELDS.get(category, [])
        snippet = json.dumps(
            parsed.get("content", [])[:20], ensure_ascii=False
        )[:3000]

        system_prompt = "你是医疗设备采购专家。严格按要求输出合法JSON，不加任何解释文字。"
        user_prompt = (
            f"文件：{parsed.get('file', '')}，类别：{category}\n"
            f"需抽取字段：{target_fields}\n"
            f"文档片段：{snippet}\n\n"
            f"输出 JSON（无法抽取填 null）：\n"
            f"{{\"fields\": {{字段名: 值}}, "
            f"\"summary\": \"50字内摘要\", "
            f"\"confidence\": 0.0-1.0}}"
        )

        t0 = time.time()
        try:
            resp = self.client.chat.completions.create(
                model=self.config.qwen_model,
                temperature=self.config.qwen_temperature,
                max_tokens=self.config.qwen_max_tokens,
                timeout=45,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user",   "content": user_prompt},
                ],
            )
            raw = resp.choices[0].message.content or "{}"
            raw = re.sub(r"^```(?:json)?\s*", "", raw.strip())
            raw = re.sub(r"\s*```$", "", raw)
            result = json.loads(raw)
        except Exception as exc:
            if repair_json:
                try:
                    result = json.loads(repair_json(raw))
                except Exception:
                    result = {}
            else:
                result = {}
            self.logger.error("qwen_extractor",
                              f"Extraction failed: {parsed.get('file', '')}: {exc}")

        self.logger.info("qwen_extractor",
                         f"Extracted: {parsed.get('file', '')}",
                         {"category": category,
                          "duration_s": round(time.time() - t0, 2),
                          "confidence": result.get("confidence")})
        return result


# ─────────────────────────────────────────────────────────────────────────────
# LangGraph state
# ─────────────────────────────────────────────────────────────────────────────
class IngestState(TypedDict, total=False):
    datasource_root: str
    output_root:     str
    config_path:     str
    app_config:      AppConfig
    logger:          IngestLogger
    parsed_docs:     Dict[str, List[Dict[str, Any]]]
    extracted:       Dict[str, List[Dict[str, Any]]]
    snapshot:        Dict[str, Any]
    exceptions:      List[str]
    started_at:      float


# ─────────────────────────────────────────────────────────────────────────────
# Workflow nodes
# ─────────────────────────────────────────────────────────────────────────────

def node_init(state: IngestState) -> IngestState:
    config = load_config(Path(state["config_path"]))
    log_path = Path(state["output_root"]) / "logs" / "ingest_runtime.jsonl"
    logger = IngestLogger(log_path)
    logger.info("init", "agent_ingest started",
                {"datasource": state["datasource_root"],
                 "parser_mode": config.parser_mode})
    return {**state, "app_config": config, "logger": logger,
            "started_at": time.time(), "exceptions": []}


def node_parse_documents(state: IngestState) -> IngestState:
    """
    A1 — parallel MinerU parsing across dual-track datasource layout.

    Dual-track scan order
    ─────────────────────
    Track A  (agent_corpus)  : <datasource>/agent_corpus/{a_requirements, …}
    Track B  (user_uploads)  : <datasource>/user_uploads/{01_requirements, …}

    Fallback: if agent_corpus/ does not exist, the legacy flat layout
    (<datasource>/a_requirements/ …) is treated as agent_corpus.

    Every parsed result is tagged with:
      • origin_type  — "agent_corpus" | "user_upload"
      • doc_category — the logical category key ("a_requirements", etc.)

    Results are merged into the same cat_key buckets so downstream nodes
    see a single unified dict keyed by logical category.
    """
    config: AppConfig    = state["app_config"]
    logger: IngestLogger = state["logger"]
    parser  = MinerUDocumentParser(config, logger)
    datasource  = Path(state["datasource_root"])
    mineru_root = Path(state["output_root"]) / "mineru_raw"

    # Initialise buckets for every known category
    parsed: Dict[str, List[Dict[str, Any]]] = {
        cat: [] for cat in CORPUS_CATEGORY_DIRS
    }

    # ── Inner helper ─────────────────────────────────────────────────────────
    def _scan_dir(
        dir_path:     Path,
        cat_key:      str,
        origin_type:  str,
        doc_category: str,
        label:        str,
    ) -> List[Dict[str, Any]]:
        """
        Parse all supported files in dir_path, tagging each result with
        origin_type and doc_category.  Returns the list of parse results.
        """
        if not dir_path.exists():
            logger.warn("parse_documents",
                        f"Directory not found, skipping: {dir_path}")
            return []

        files = [
            f for f in dir_path.iterdir()
            if f.suffix.lower() in SUPPORTED_EXTENSIONS
            and not f.name.startswith("~$")
        ]
        logger.info(
            "parse_documents",
            f"Scanning {label} [{origin_type}]: {len(files)} file(s)",
        )
        results: List[Dict[str, Any]] = []
        slug = f"{origin_type[:6]}_{cat_key}"          # unique MinerU output sub-dir

        with ThreadPoolExecutor(max_workers=config.max_workers) as pool:
            future_map = {
                pool.submit(parser.parse, f, mineru_root / slug / f.stem): f
                for f in files
            }
            for future in as_completed(future_map):
                fp = future_map[future]
                try:
                    r = future.result()
                    r["category"]     = cat_key
                    r["origin_type"]  = origin_type
                    r["doc_category"] = doc_category
                    results.append(r)
                except Exception as exc:
                    state["exceptions"].append(f"parse:{fp.name}:{exc}")
                    results.append({
                        "file":         fp.name,
                        "category":     cat_key,
                        "origin_type":  origin_type,
                        "doc_category": doc_category,
                        "content":      [],
                        "error":        str(exc),
                    })
        return results

    # ── Track A: agent_corpus ─────────────────────────────────────────────────
    corpus_root = datasource / CORPUS_ROOT
    if corpus_root.exists():
        logger.info("parse_documents",
                    f"Track A (agent_corpus): scanning {corpus_root}")
        for cat_key, cat_label in CORPUS_CATEGORY_DIRS.items():
            parsed[cat_key] += _scan_dir(
                corpus_root / cat_key,
                cat_key,
                "agent_corpus",
                cat_key,
                f"{CORPUS_ROOT}/{cat_key} ({cat_label})",
            )
    else:
        # Fallback: legacy flat layout — treat as agent_corpus
        logger.warn(
            "parse_documents",
            f"{CORPUS_ROOT}/ not found under {datasource} — "
            "falling back to legacy flat layout (treated as agent_corpus)",
        )
        for cat_key, cat_label in CORPUS_CATEGORY_DIRS.items():
            parsed[cat_key] += _scan_dir(
                datasource / cat_key,
                cat_key,
                "agent_corpus",
                cat_key,
                f"{cat_key} ({cat_label}) [legacy-flat]",
            )

    # ── Track B: user_uploads ─────────────────────────────────────────────────
    uploads_root = datasource / UPLOADS_ROOT
    if uploads_root.exists():
        logger.info("parse_documents",
                    f"Track B (user_uploads): scanning {uploads_root}")
        for upload_dir, cat_key in UPLOAD_DIR_TO_CATEGORY.items():
            parsed[cat_key] += _scan_dir(
                uploads_root / upload_dir,
                cat_key,
                "user_upload",
                cat_key,
                f"{UPLOADS_ROOT}/{upload_dir} → {cat_key}",
            )
    else:
        logger.info(
            "parse_documents",
            f"{UPLOADS_ROOT}/ not found under {datasource} — "
            "no user-upload data this run (corpus-only mode)",
        )

    _save_json(
        Path(state["output_root"]) / "intermediate" / "parsed_docs.json",
        parsed,
    )
    return {**state, "parsed_docs": parsed}


def node_extract_fields(state: IngestState) -> IngestState:
    """
    A2 — parallel Qwen field extraction per document.

    Propagates origin_type and doc_category from each parsed doc into the
    extracted item so that node_build_snapshot can apply two-pass priority.
    """
    config: AppConfig    = state["app_config"]
    logger: IngestLogger = state["logger"]
    extractor = QwenFieldExtractor(config, logger)
    parsed    = state.get("parsed_docs", {})
    extracted: Dict[str, List[Dict[str, Any]]] = {}

    for cat_dir, doc_list in parsed.items():
        cat_results: List[Dict[str, Any]] = []
        valid_docs = [d for d in doc_list if not d.get("error") and not d.get("skipped")]

        with ThreadPoolExecutor(max_workers=config.max_workers) as pool:
            future_map = {
                pool.submit(extractor.extract, cat_dir, d): d
                for d in valid_docs
            }
            for future in as_completed(future_map):
                doc = future_map[future]
                try:
                    result = future.result()
                    cat_results.append({
                        "file":         doc.get("file", ""),
                        "category":     cat_dir,
                        # ── dual-track tags propagated from parsed doc ──────
                        "origin_type":  doc.get("origin_type",  "agent_corpus"),
                        "doc_category": doc.get("doc_category", cat_dir),
                        # ── extraction results ──────────────────────────────
                        "fields":       result.get("fields", {}),
                        "summary":      result.get("summary", ""),
                        "confidence":   result.get("confidence"),
                    })
                except Exception as exc:
                    state["exceptions"].append(f"extract:{doc.get('file','')}:{exc}")

        extracted[cat_dir] = cat_results

    _save_json(
        Path(state["output_root"]) / "intermediate" / "extracted_fields.json",
        extracted,
    )
    logger.info("extract_fields", "Field extraction complete",
                {k: len(v) for k, v in extracted.items()})
    return {**state, "extracted": extracted}


def node_build_snapshot(state: IngestState) -> IngestState:
    """
    Merge extracted fields into project_snapshot.json — dual-track edition.

    Two-pass priority strategy
    ──────────────────────────
    Within each category bucket, items are sorted so user_upload items are
    processed before agent_corpus items.  Scalar fields (device_name, charge
    rate, utilisation, etc.) are stored as ConflictableField objects that
    carry the winning value plus a full audit trail.  The priority rules are
    enforced by _make_conflictable_field / _merge_conflictable_field:

      • user_upload always wins over agent_corpus for .value
      • When the same metric arrives from both tracks with different numeric
        values, conflict is set to True and all_values[] preserves both
      • Two user_upload docs with different values also set conflict = True

    Backward compatibility
    ──────────────────────
    The top-level revenue_params keys charge_conflict and charge_per_exam_all
    are computed *after* the two-pass loop and mirror the ConflictableField
    data so that agent_approval.py can read either schema without change.

    Retained fixes
    ──────────────
    FIX-1  Non-MR documents excluded from canonical key_params
    FIX-2  Cross-category competitor data in a_requirements harvested
    FIX-4  B类内部/未广审宣传材料、1.5T竞品隔离
    FIX-5b Missing price detection → missing_price_competitors
    """
    logger    = state["logger"]
    extracted = state.get("extracted", {})

    snapshot: Dict[str, Any] = {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "parser_mode":  state["app_config"].parser_mode,
        "categories":   extracted,
        "key_params": {
            # ConflictableField or None for scalar identity params
            "device_name":         None,
            "applicant_dept":      None,
            "request_qty":         None,
            # Lists — each entry carries origin_type + doc_category tags
            "competitor_list":     [],
            "registration_certs":  [],
            # ConflictableField or None
            "daily_utilization":   None,
            "maintenance_records": [],
            "revenue_params": {
                # ── ConflictableField schema ──────────────────────────────
                "charge_per_exam": None,   # {value, origin_type, source_file,
                "monthly_volume":  None,   #  doc_category, conflict, all_values}
                "payback_period":  None,
                # ── Backward-compat flat fields (derived after two-pass) ──
                "charge_conflict":     False,
                "charge_per_exam_all": [],
            },
        },
        "excluded_non_mr":            [],
        "excluded_internal_promo":    [],
        "reference_only_competitors": [],
        "missing_price_competitors":  [],
    }

    _MISSING_PRICE = {"", "null", "none", "——", "—", "n/a", "无", "待询价", "tbd"}

    # ── Two-pass ordering: user_upload items first, agent_corpus second ───────
    def _priority_key(item: Dict[str, Any]) -> int:
        return 0 if item.get("origin_type") == "user_upload" else 1

    for cat, items in extracted.items():
        ordered = sorted(items, key=_priority_key)

        for item in ordered:
            f            = item.get("fields", {})
            origin_type  = item.get("origin_type",  "agent_corpus")
            doc_category = item.get("doc_category", cat)
            src_file     = item.get("file", "")

            # ── a_requirements ────────────────────────────────────────────
            if cat == "a_requirements":

                # FIX-1: detect and skip non-MR documents for canonical fields
                if _is_non_mr_document(f):
                    snapshot["excluded_non_mr"].append({
                        "file":        src_file,
                        "device_name": f.get("设备名称"),
                        "device_type": f.get("设备类型"),
                        "origin_type": origin_type,
                        "reason":      "non-MR device detected — excluded from key_params",
                    })
                    logger.warn(
                        "build_snapshot",
                        f"Non-MR doc excluded from key_params: {src_file}",
                        {"device": f.get("设备名称"), "type": f.get("设备类型"),
                         "origin_type": origin_type},
                    )
                    continue

                kp = snapshot["key_params"]

                # ── ConflictableField: device identity scalars ─────────────
                for field_key, raw_key in [
                    ("device_name",    "设备名称"),
                    ("applicant_dept", "申请科室"),
                    ("request_qty",    "申请数量"),
                ]:
                    raw_val = f.get(raw_key)
                    if raw_val:
                        if kp[field_key] is None:
                            kp[field_key] = _make_conflictable_field(
                                raw_val, origin_type, src_file, doc_category)
                        else:
                            _merge_conflictable_field(
                                kp[field_key], raw_val,
                                origin_type, src_file, doc_category)

                # FIX-2: cross-category competitor data embedded in A docs
                competitor_raw = f.get("竞品品牌列表") or f.get("竞品参数对比")
                if competitor_raw:
                    kp["competitor_list"].append({
                        "brand":        competitor_raw,
                        "model":        None,
                        "price":        None,
                        "source":       src_file,
                        "origin":       "a_requirements",
                        "origin_type":  origin_type,
                        "doc_category": doc_category,
                    })
                    logger.info(
                        "build_snapshot",
                        f"Cross-category competitor data in a_requirements: {src_file}",
                        {"origin_type": origin_type},
                    )

                # ── Revenue params as ConflictableFields ───────────────────
                rp = kp["revenue_params"]
                charge_raw = f.get("单次收费标准")
                if charge_raw:
                    if rp["charge_per_exam"] is None:
                        rp["charge_per_exam"] = _make_conflictable_field(
                            charge_raw, origin_type, src_file, doc_category)
                    else:
                        _merge_conflictable_field(
                            rp["charge_per_exam"], charge_raw,
                            origin_type, src_file, doc_category)

                for field_key, raw_key in [
                    ("monthly_volume", "月检查量"),
                    ("payback_period", "投入回收期"),
                ]:
                    raw_val = f.get(raw_key)
                    if raw_val:
                        if rp[field_key] is None:
                            rp[field_key] = _make_conflictable_field(
                                raw_val, origin_type, src_file, doc_category)
                        else:
                            _merge_conflictable_field(
                                rp[field_key], raw_val,
                                origin_type, src_file, doc_category)

            # ── b_competitors ─────────────────────────────────────────────
            elif cat == "b_competitors":
                if _is_internal_use_item(item):
                    snapshot["excluded_internal_promo"].append({
                        "file":        src_file,
                        "brand":       f.get("品牌"),
                        "origin_type": origin_type,
                        "reason":      "internal / not ad-review approved",
                    })
                    logger.warn(
                        "build_snapshot",
                        f"B类内部或未广审材料已排除: {src_file}",
                        {"brand": f.get("品牌"), "origin_type": origin_type},
                    )
                    continue

                entry = {
                    "brand":          f.get("品牌"),
                    "model":          f.get("型号"),
                    "price":          f.get("含税报价"),
                    "magnetic_field": f.get("磁场强度"),
                    "source":         src_file,
                    "origin":         "b_competitors",
                    # dual-track tags
                    "origin_type":    origin_type,
                    "doc_category":   doc_category,
                }

                if _item_is_low_field_15t(
                    f, str(src_file), str(item.get("summary") or "")
                ):
                    snapshot["reference_only_competitors"].append({
                        **entry,
                        "reason": "1.5T (or low-field) product — excluded from 3.0T comparison set",
                    })
                    logger.info(
                        "build_snapshot",
                        f"Low-field (≈1.5T) competitor → reference_only: {src_file}",
                        {"brand": f.get("品牌"), "origin_type": origin_type},
                    )
                    continue

                if not f.get("品牌"):
                    continue

                snapshot["key_params"]["competitor_list"].append(entry)

                # FIX-5b: flag missing / invalid price
                _price_raw = str(f.get("含税报价") or "").strip()
                if _price_raw.lower() in _MISSING_PRICE:
                    snapshot["missing_price_competitors"].append({
                        "brand":       f.get("品牌"),
                        "model":       f.get("型号"),
                        "source":      src_file,
                        "origin_type": origin_type,
                        "reason":      "含税报价字段为空或无效值",
                    })
                    logger.warn(
                        "build_snapshot",
                        f"竞品含税报价缺失: {f.get('品牌')} {f.get('型号')}",
                        {"file": src_file, "origin_type": origin_type},
                    )

            # ── c_compliance ──────────────────────────────────────────────
            elif cat == "c_compliance" and f.get("注册证号"):
                snapshot["key_params"]["registration_certs"].append({
                    "cert_no":      f.get("注册证号"),
                    "device":       f.get("设备名称"),
                    "source":       src_file,
                    "origin_type":  origin_type,
                    "doc_category": doc_category,
                })

            # ── d_operations ──────────────────────────────────────────────
            elif cat == "d_operations":
                kp = snapshot["key_params"]
                raw_util = f.get("工作饱和度")
                if raw_util:
                    if kp["daily_utilization"] is None:
                        kp["daily_utilization"] = _make_conflictable_field(
                            raw_util, origin_type, src_file, doc_category)
                    else:
                        _merge_conflictable_field(
                            kp["daily_utilization"], raw_util,
                            origin_type, src_file, doc_category)

    # ── Backward-compat flat fields derived from ConflictableField ────────────
    _rp = snapshot["key_params"]["revenue_params"]
    _ce = _rp["charge_per_exam"]
    if _ce is not None:
        _rp["charge_conflict"]     = _ce.get("conflict", False)
        _rp["charge_per_exam_all"] = _ce.get("all_values", [])
        if _rp["charge_conflict"]:
            logger.warn(
                "build_snapshot",
                f"收费标准冲突 (dual-track)：conflict=True，"
                f"{len(_rp['charge_per_exam_all'])} 个来源",
                {"all_values": _rp["charge_per_exam_all"]},
            )

    # ── Persist ───────────────────────────────────────────────────────────────
    results_dir = Path(state["output_root"]) / "results"
    results_dir.mkdir(parents=True, exist_ok=True)
    _save_json(results_dir / "project_snapshot.json", snapshot)

    _ce_val = _ce.get("value") if _ce else None
    logger.info("build_snapshot", "project_snapshot.json written (dual-track)", {
        "competitors":                len(snapshot["key_params"]["competitor_list"]),
        "reference_only_competitors": len(snapshot["reference_only_competitors"]),
        "excluded_internal_promo":    len(snapshot["excluded_internal_promo"]),
        "certs":                      len(snapshot["key_params"]["registration_certs"]),
        "device":                     _rp_device(snapshot),
        "excluded_non_mr":            len(snapshot["excluded_non_mr"]),
        "charge_per_exam":            _ce_val,
        "charge_conflict":            _rp["charge_conflict"],
        "charge_sources":             len(_rp["charge_per_exam_all"]),
        "missing_price_competitors":  len(snapshot["missing_price_competitors"]),
    })
    return {**state, "snapshot": snapshot}


def _rp_device(snapshot: Dict[str, Any]) -> Optional[str]:
    """Extract resolved device_name from ConflictableField or plain string."""
    dn = snapshot.get("key_params", {}).get("device_name")
    if dn is None:
        return None
    if isinstance(dn, dict):
        return dn.get("value")
    return str(dn)


def node_export_log(state: IngestState) -> IngestState:
    """Export a PDF run report (ReportLab — same pattern as supercare)."""
    logger     = state["logger"]
    output_root = Path(state["output_root"])
    pdf_path    = output_root / "logs" / "ingest_report.pdf"

    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        doc    = SimpleDocTemplate(str(pdf_path), pagesize=A4)
        styles = getSampleStyleSheet()
        snap   = state.get("snapshot", {}).get("key_params", {})
        story  = [
            Paragraph("agent_ingest — Run Report", styles["Title"]),
            Spacer(1, 12),
        ]
        lines = [
            f"Generated  : {state.get('snapshot', {}).get('generated_at', '')}",
            f"Parser mode: {state['app_config'].parser_mode}",
            f"Device     : {snap.get('device_name', 'n/a')}",
            f"Dept       : {snap.get('applicant_dept', 'n/a')}",
            f"Competitors: {len(snap.get('competitor_list', []))}",
            f"Certs      : {len(snap.get('registration_certs', []))}",
            f"Exceptions : {len(state.get('exceptions', []))}",
            f"Duration   : {round(time.time() - state.get('started_at', time.time()), 1)}s",
        ]
        for line in lines:
            story.append(Paragraph(line, styles["Normal"]))
            story.append(Spacer(1, 6))
        doc.build(story)
        logger.info("export_log", "PDF report written", {"path": str(pdf_path)})
    except Exception as exc:
        logger.error("export_log", f"PDF export failed: {exc}")

    return state


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────
def _save_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as fh:
        json.dump(data, fh, ensure_ascii=False, indent=2)


# ─────────────────────────────────────────────────────────────────────────────
# Build workflow
# ─────────────────────────────────────────────────────────────────────────────
def build_workflow():
    g = StateGraph(IngestState)
    g.add_node("init",             node_init)
    g.add_node("parse_documents",  node_parse_documents)
    g.add_node("extract_fields",   node_extract_fields)
    g.add_node("build_snapshot",   node_build_snapshot)
    g.add_node("export_log",       node_export_log)

    g.add_edge(START,            "init")
    g.add_edge("init",           "parse_documents")
    g.add_edge("parse_documents","extract_fields")
    g.add_edge("extract_fields", "build_snapshot")
    g.add_edge("build_snapshot", "export_log")
    g.add_edge("export_log",     END)
    return g.compile()


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────
def main() -> None:
    parser = argparse.ArgumentParser(
        description="agent_ingest — document parsing & field extraction")
    parser.add_argument("--datasource",  default=DEFAULT_DATASOURCE)
    parser.add_argument("--output-root", default=DEFAULT_OUTPUT_ROOT)
    parser.add_argument("--config",      default=DEFAULT_CONFIG_PATH)
    parser.add_argument("--parser-mode", default="mineru",
                        choices=["mineru", "openpyxl"],
                        help="mineru = full OCR pipeline; openpyxl = xlsx-only baseline")
    args = parser.parse_args()

    wf = build_workflow()
    final = wf.invoke({
        "datasource_root": args.datasource,
        "output_root":     args.output_root,
        "config_path":     args.config,
    })

    exceptions = final.get("exceptions", [])
    print(f"\n✅  agent_ingest complete  |  exceptions: {len(exceptions)}")
    print(f"    project_snapshot  →  {args.output_root}/results/project_snapshot.json")
    print(f"    PDF report        →  {args.output_root}/logs/ingest_report.pdf")
    if exceptions:
        print("⚠   exception log:")
        for e in exceptions:
            print(f"    · {e}")


if __name__ == "__main__":
    main()
