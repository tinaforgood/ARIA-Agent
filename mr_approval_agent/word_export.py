"""
word_export.py — 将审批流水线 JSON + Markdown 正文导出为专业 Word（.docx）

改进版 v2：
  1. 专业封面页（医院名称、设备名称、版本、日期、编制信息）
  2. Markdown 表格 → 真实 Word 表格（含表头加粗/居中）
  3. **bold** / *italic* / `code` → Word 格式化 runs
  4. --- 水平线 → 浅色上边框空段落
  5. > 引用块 → 带蓝色左边框的提示段落
  6. 移除原始 JSON 数据块，改为精简结构化附录
  7. 修复标题层级（# → Heading 1, ## → Heading 2, ### → Heading 3）
  8. 首行缩进（中文习惯 2 字符）
  9. 从 results_dir 向上读取 metadata.json 获取 hospital_name
"""

from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path
from typing import Any, List, Mapping, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ── 字体辅助 ──────────────────────────────────────────────────────────────────

def _set_east_asia(run, zh_font: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), zh_font)


def _apply_font(
    run,
    zh_font: str = "宋体",
    size_pt: Optional[float] = None,
    bold: bool = False,
    italic: bool = False,
    color: Optional[RGBColor] = None,
) -> None:
    _set_east_asia(run, zh_font)
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color


# ── 内联 Markdown 解析（**bold**, *italic*, `code`）────────────────────────────

_INLINE_RE = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`')


def _add_inline(
    para,
    text: str,
    zh_font: str = "宋体",
    size_pt: Optional[float] = None,
    base_bold: bool = False,
    base_color: Optional[RGBColor] = None,
) -> None:
    """将含 **bold**/*italic*/`code` 的文本以格式化 runs 写入段落。"""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            _apply_font(run, zh_font, size_pt, bold=base_bold, color=base_color)
        if m.group(1) is not None:          # **bold**
            run = para.add_run(m.group(1))
            _apply_font(run, zh_font, size_pt, bold=True, color=base_color)
        elif m.group(2) is not None:        # *italic*
            run = para.add_run(m.group(2))
            _apply_font(run, zh_font, size_pt, italic=True, color=base_color)
        elif m.group(3) is not None:        # `code`
            run = para.add_run(m.group(3))
            run.font.name = "Courier New"
            if size_pt:
                run.font.size = Pt(size_pt - 0.5)
        pos = m.end()
    if pos < len(text):
        run = para.add_run(text[pos:])
        _apply_font(run, zh_font, size_pt, bold=base_bold, color=base_color)


# ── 标题辅助 ──────────────────────────────────────────────────────────────────

_HEADING_CFG = {
    1: ("黑体", 18, True),
    2: ("黑体", 15, True),
    3: ("黑体", 13, True),
    4: ("宋体", 12, True),
}


def _add_heading(doc: Document, text: str, level: int) -> None:
    level = max(1, min(level, 4))
    h = doc.add_heading("", level=level)
    h.clear()
    zh_font, pt, bold = _HEADING_CFG[level]
    run = h.add_run(text)
    _apply_font(run, zh_font, pt, bold=bold)
    h.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    h.paragraph_format.space_after = Pt(4)


# ── 段落边框辅助 ──────────────────────────────────────────────────────────────

def _add_left_border(para, color_hex: str = "4472C4", sz: str = "12", space: str = "8") -> None:
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_top_border(para, color_hex: str = "CCCCCC", sz: str = "6") -> None:
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), sz)
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color_hex)
    pBdr.append(top)
    pPr.append(pBdr)


# ── Markdown 表格 ─────────────────────────────────────────────────────────────

def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and len(s) > 2


def _is_separator_row(line: str) -> bool:
    return _is_table_row(line) and re.match(r"^[\|\s\-:]+$", line.strip())


def _parse_md_table(lines: List[str]) -> List[List[str]]:
    rows: List[List[str]] = []
    for ln in lines:
        if _is_separator_row(ln):
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def _add_word_table(doc: Document, md_lines: List[str]) -> None:
    rows = _parse_md_table(md_lines)
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    rows = [r + [""] * (num_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for i, row in enumerate(rows):
        is_header = (i == 0)
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.paragraphs[0].clear()
            _add_inline(
                cell.paragraphs[0],
                cell_text,
                zh_font="黑体" if is_header else "宋体",
                size_pt=9.5,
                base_bold=is_header,
            )
            if is_header:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)


# ── 专业封面页 ────────────────────────────────────────────────────────────────

def _add_cover_page(doc: Document, state: Mapping[str, Any]) -> None:
    hospital = (state.get("hospital_name") or "").strip()
    req = state.get("requirements_result") or {}
    device_name = (req.get("device_name") or "医疗设备").strip()
    dept = (req.get("department") or "").strip()
    today_str = date.today().strftime("%Y年%m月%d日")

    # 垂直留白
    for _ in range(6):
        doc.add_paragraph()

    # 主标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(8)
    run = title_p.add_run("医疗设备立项建议书")
    _apply_font(run, "黑体", 28, bold=True)

    # 副标题（设备名）
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(30)
    run = sub_p.add_run(f"——{device_name}采购项目")
    _apply_font(run, "宋体", 16)

    # 分隔线
    sep = doc.add_paragraph("─" * 28)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_after = Pt(24)
    for r in sep.runs:
        r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # 元数据信息
    def _meta_line(label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        r1 = p.add_run(f"{label}：")
        _apply_font(r1, "黑体", 12, bold=True)
        r2 = p.add_run(value)
        _apply_font(r2, "宋体", 12)

    if hospital:
        _meta_line("申报单位", hospital)
    if dept:
        _meta_line("申报科室", dept)
    _meta_line("文件版本", "V1.0（AI 辅助生成，供内部审批参考）")
    _meta_line("编制日期", today_str)
    _meta_line("文件性质", "内部审批文件  ·  保密级别：内部")

    doc.add_page_break()


# ── Markdown → Word 正文渲染 ──────────────────────────────────────────────────

def _render_markdown(doc: Document, md_text: str) -> None:
    """将 agent6 产出的 Markdown 文本转为 Word 段落/标题/表格。"""
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        # ── 空行
        if not stripped:
            i += 1
            continue

        # ── 水平分割线 --- / === / ***
        if re.match(r"^[-=\*]{3,}$", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            _add_top_border(p)
            i += 1
            continue

        # ── 标题 ####
        if line.startswith("#### "):
            _add_heading(doc, line[5:].strip(), level=4)
            i += 1
            continue
        # ── 标题 ###
        if line.startswith("### "):
            _add_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue
        # ── 标题 ##
        if line.startswith("## "):
            _add_heading(doc, line[3:].strip(), level=2)
            i += 1
            continue
        # ── 标题 #（一级：跳过重复的文档总标题）
        if line.startswith("# "):
            _add_heading(doc, line[2:].strip(), level=1)
            i += 1
            continue

        # ── Markdown 表格（收集连续表格行）
        if _is_table_row(stripped):
            table_lines: List[str] = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                table_lines.append(lines[i])
                i += 1
            _add_word_table(doc, table_lines)
            continue

        # ── 无序列表
        m_bullet = re.match(r"^[\-\*\+]\s+(.+)$", stripped)
        if m_bullet:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(2)
            _add_inline(p, m_bullet.group(1), "宋体", 10.5)
            i += 1
            continue

        # ── 有序列表
        m_num = re.match(r"^\d+[\.)]\s+(.+)$", stripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(2)
            _add_inline(p, m_num.group(1), "宋体", 10.5)
            i += 1
            continue

        # ── 引用块 >
        if stripped.startswith(">"):
            content = re.sub(r"^>+\s*", "", stripped)
            # 连续引用行合并
            combined = content
            while i + 1 < len(lines) and lines[i + 1].strip().startswith(">"):
                i += 1
                combined += " " + re.sub(r"^>+\s*", "", lines[i].strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            _add_left_border(p, color_hex="4472C4", sz="12", space="8")
            _add_inline(p, combined, "宋体", 10,
                        base_color=RGBColor(0x44, 0x44, 0x66))
            i += 1
            continue

        # ── 普通段落（首行缩进 2 字符）
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(21)   # 10.5pt × 2
        p.paragraph_format.space_after = Pt(4)
        _add_inline(p, stripped, "宋体", 10.5)
        i += 1


# ── 精简结构化附录 ─────────────────────────────────────────────────────────────

def _add_appendix(doc: Document, state: Mapping[str, Any]) -> None:
    """在正文末尾追加精简附录：质量评审得分 + 关键参数。"""
    feedback = state.get("feedback_result") or {}
    if not feedback:
        return

    doc.add_page_break()
    _add_heading(doc, "附录  AI 质量评审摘要", level=1)

    scores = feedback.get("scores") or {}
    if scores:
        # 分数概览表
        score_rows = [["评审维度", "得分（/10）"]]
        labels = {
            "completeness": "材料完整性",
            "data_accuracy": "数据准确性",
            "compliance_coverage": "合规覆盖度",
            "financial_rigor": "财务严谨性",
            "writing_quality": "文字质量",
        }
        for key, label in labels.items():
            v = scores.get(key)
            if v is not None:
                score_rows.append([label, str(v)])
        total = feedback.get("total_score")
        grade = feedback.get("grade", "")
        if total is not None:
            score_rows.append(["综合评分", f"{total}  ({grade})"])

        # 转为 Word 表格
        table = doc.add_table(rows=len(score_rows), cols=2)
        table.style = "Table Grid"
        for i, (c0, c1) in enumerate(score_rows):
            for j, text in enumerate([c0, c1]):
                cell = table.rows[i].cells[j]
                cell.paragraphs[0].clear()
                is_header = (i == 0)
                run = cell.paragraphs[0].add_run(text)
                _apply_font(run, "黑体" if is_header else "宋体", 10, bold=is_header)
                if is_header:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

    # 建议摘要
    rec = feedback.get("approval_recommendation") or ""
    missing = feedback.get("missing_items") or []
    if rec:
        p = doc.add_paragraph()
        r1 = p.add_run("审批建议：")
        _apply_font(r1, "黑体", 11, bold=True)
        r2 = p.add_run(rec)
        _apply_font(r2, "宋体", 11)
    if missing:
        p2 = doc.add_paragraph()
        r = p2.add_run("待补充材料：" + "、".join(str(m) for m in missing))
        _apply_font(r, "宋体", 10.5)


# ── 从 results_dir 向上找 metadata.json ──────────────────────────────────────

def _read_hospital_from_metadata(results_dir: Path) -> str:
    """
    results_dir = cases/{id}/outputs/approval/results
    metadata    = cases/{id}/metadata.json
    """
    try:
        meta_path = results_dir.parent.parent.parent.parent / "metadata.json"
        if meta_path.is_file():
            meta = json.loads(meta_path.read_text(encoding="utf-8"))
            return meta.get("hospital_name") or ""
    except Exception:
        pass
    return ""


# ── 主导出函数 ────────────────────────────────────────────────────────────────

def default_template_path() -> Path:
    return Path(__file__).resolve().parent / "templates" / "approval_proposal_template.docx"


def export_approval_word(
    state: Mapping[str, Any],
    results_dir: Path,
    *,
    out_name: str = "project_document.docx",
) -> Path:
    """
    生成专业格式 Word 立项建议书。
    - 封面页（医院/设备/日期）
    - 正文（agent6 Markdown → 真实 Word 元素）
    - 附录（AI 质量评审摘要）
    返回写入路径；异常向上抛出，由调用方记录日志。
    """
    results_dir.mkdir(parents=True, exist_ok=True)
    out_path = results_dir / out_name

    # 加载模板仅为获取样式，清除占位内容
    tpl = default_template_path()
    if tpl.is_file():
        doc = Document(str(tpl))
        # 清除模板中的所有占位段落和表格
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
        for t in list(doc.tables):
            t._element.getparent().remove(t._element)
    else:
        doc = Document()

    # 尝试从 state 或 metadata.json 获取 hospital_name
    hospital_name = (
        state.get("hospital_name")
        or _read_hospital_from_metadata(results_dir)
        or ""
    )
    # 注入到 state（不可变 Mapping 时用局部 dict）
    enriched: dict = dict(state)
    enriched.setdefault("hospital_name", hospital_name)

    # ── 1. 封面页
    _add_cover_page(doc, enriched)

    # ── 2. 正文（agent6 Markdown）
    md_text = str(state.get("project_document") or "").strip()
    # 去除开头的文档总标题块（已由封面页呈现，避免重复）
    # 策略：找到第一条水平分割线（---）并从其后开始；
    # 若无分割线，则只跳过最开头的 # / ## 行（不超过3行）
    md_lines_stripped = md_text.splitlines()
    skip_until = 0
    first_sep = None
    for idx, ln in enumerate(md_lines_stripped):
        s = ln.strip()
        if re.match(r"^[-=\*]{3,}$", s):
            first_sep = idx
            break
    if first_sep is not None:
        # 跳过分割线及之前的所有行（标题块），从分割线之后开始
        skip_until = first_sep + 1
    else:
        # 无分割线：跳过最多前3行中的标题/空行
        for idx, ln in enumerate(md_lines_stripped[:3]):
            s = ln.strip()
            if s.startswith("#") or not s:
                skip_until = idx + 1
            else:
                break
    md_text = "\n".join(md_lines_stripped[skip_until:])
    if md_text:
        _render_markdown(doc, md_text)
    else:
        p = doc.add_paragraph("（正文为空，请检查 Agent6 输出）")
        _apply_font(p.runs[0] if p.runs else p.add_run(""), "宋体", 11)

    # ── 3. 附录：AI 质量评审摘要
    _add_appendix(doc, state)

    doc.save(str(out_path))
    return out_path
